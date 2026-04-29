#!/usr/bin/env python3
"""Generate Claude-VAI.docx — full deployment workflow walkthrough."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1a73e8")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "f0f4ff")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


def add_code_block(doc, text):
    for line in text.strip().split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        # Set shading on the paragraph
        pPr = p._element.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {
            qn("w:fill"): "F5F5F5",
            qn("w:val"): "clear",
        })
        pPr.append(shd)


def add_diagram(doc, text):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.5)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def build_document():
    doc = Document()

    # -- Styles --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    for level in range(1, 5):
        sname = f"Heading {level}"
        s = doc.styles[sname]
        s.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
        if level == 1:
            s.font.size = Pt(22)
        elif level == 2:
            s.font.size = Pt(16)
        elif level == 3:
            s.font.size = Pt(13)
        else:
            s.font.size = Pt(11)

    # ===== TITLE PAGE =====
    for _ in range(6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run("Claude Code on Vertex AI")
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Full Deployment & Operations Walkthrough")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x5F, 0x63, 0x68)

    for _ in range(2):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Google Cloud  |  Anthropic  |  Reference Architecture v1.0")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

    doc.add_page_break()

    # ===== TABLE OF CONTENTS =====
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        ("", "Introduction and Background"),
        ("", "What You'll Get"),
        ("", "Deployment Modes"),
        ("", "Prerequisites and Deployment Paths"),
        ("1", "The Orchestrator: deploy.sh"),
        ("2", "The Core Gateway: deploy-llm-gateway.sh"),
        ("3", "Shared MCP Tools: deploy-mcp-gateway.sh"),
        ("4", "Self-Service Portal: deploy-dev-portal.sh"),
        ("5", "Logging Pipeline: deploy-observability.sh"),
        ("6", "Runtime Request Flow"),
        ("7", "Developer Setup: developer-setup.sh"),
        ("8", "MCP Gateway Runtime"),
        ("9", "Populating the Dashboard: seed-demo-data.sh"),
        ("10", "Viewing the Admin Dashboard"),
        ("11", "Access Tiers and Security Model"),
        ("12", "Developer Onboarding"),
        ("13", "Validation and Regression Testing"),
        ("A", "Complete Lifecycle Diagram"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        label = f"  {num}.  {title}" if num else f"  {title}"
        run = p.add_run(label)
        run.font.size = Pt(11)
    doc.add_page_break()

    # =================================================================
    # INTRODUCTION AND BACKGROUND
    # =================================================================
    doc.add_heading("Introduction and Background", level=1)

    doc.add_paragraph(
        "This document is a full deployment and operations walkthrough for "
        "Claude Code on GCP via Vertex AI — a production-quality reference "
        "architecture and deployment kit for running Claude Code on Google "
        "Cloud with all model inference routed through Vertex AI."
    )
    doc.add_paragraph(
        "No traffic to api.anthropic.com. Google identity everywhere. "
        "Near-zero cost when idle. Built for teams whose security reviewers "
        "will actually read the diagram."
    )

    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "This repository is a reference architecture, not a supported product "
        "of Google LLC, Anthropic PBC, or any affiliated entity. It is released "
        "under Apache 2.0 and is provided strictly as-is, with no warranty of "
        "any kind, express or implied."
    )
    doc.add_paragraph(
        "You are responsible for reviewing, adapting, and testing every component "
        "before running it against a production Google Cloud project. The "
        "operational, security, compliance, and cost consequences of any "
        "deployment are yours to evaluate and own."
    )
    p = doc.add_paragraph()
    run = p.add_run("Use at your own risk — this is a map, not the territory.")
    run.bold = True

    # =================================================================
    # WHAT YOU'LL GET
    # =================================================================
    doc.add_heading("What You'll Get", level=1)

    doc.add_paragraph(
        "When deployment completes, your GCP project will contain the "
        "following components:"
    )
    add_table(doc,
        ["Component", "What It Is", "Why It Exists"],
        [
            ["LLM Gateway", "Cloud Run service — tiny FastAPI reverse proxy",
             "Single point for auth, logging, and header sanitation in front of Vertex AI"],
            ["MCP Gateway", "Cloud Run service — FastMCP over Streamable HTTP",
             "Place to host your organization's custom MCP tools"],
            ["Dev Portal", "Cloud Run static site",
             "Self-service setup instructions for developers (IAP-protected in GLB mode, Cloud Run IAM in standard mode)"],
            ["Dev VM (optional)", "GCE VM with VS Code Server, accessed via IAP",
             "Cloud dev environment for teams that don't want local installs"],
            ["Observability", "Log sink to BigQuery + built-in admin dashboard",
             "Admin dashboard: who is using Claude, how much, errors, top models"],
        ],
    )

    doc.add_paragraph("")
    doc.add_heading("Authentication Model", level=2)
    doc.add_paragraph(
        "In standard mode, the LLM and MCP gateways use app-level token "
        "validation (token_validation.py middleware) that accepts both OAuth2 "
        "access tokens and OIDC identity tokens. Cloud Run's built-in invoker "
        "IAM check is disabled (--no-invoker-iam-check) because Claude Code "
        "sends access tokens, which Cloud Run IAM rejects. An ALLOWED_PRINCIPALS "
        "allowlist controls who can call the gateways."
    )

    # =================================================================
    # DEPLOYMENT MODES
    # =================================================================
    doc.add_heading("Deployment Modes", level=1)

    doc.add_paragraph(
        "The deployment kit supports three mutually exclusive ingress modes. "
        "Each mode is a single deploy.sh prompt — no manual Cloud Run "
        "configuration needed."
    )

    doc.add_heading("Standard Mode (default)", level=2)
    doc.add_paragraph(
        "Cloud Run services use --ingress all so developer laptops can reach "
        "them directly. The token_validation.py middleware is the security "
        "boundary. This is the simplest mode and suitable for initial "
        "deployments and small teams."
    )

    doc.add_heading("GLB Mode (optional)", level=2)
    doc.add_paragraph(
        "A Global HTTP(S) Load Balancer sits in front of all services. "
        "Ingress is internal-and-cloud-load-balancing; the GLB is the only "
        "entry point. Gateways use the same app-level token validation as "
        "standard mode; portal and dashboard use IAP for browser auth. "
        "Requires a DNS domain for managed certificates."
    )
    add_diagram(doc,
        "Developer laptop --(HTTPS + access token)--> GLB --> Cloud Run (internal + GLB)\n"
        "                 --(browser + IAP SSO)------> GLB --> Cloud Run (IAP-protected)"
    )

    doc.add_heading("VPC Internal Mode (optional, mutually exclusive with GLB)", level=2)
    doc.add_paragraph(
        "All Cloud Run services use --ingress internal — they are only "
        "reachable from within the VPC. Developers access services through "
        "the dev VM, which is inside the VPC and accessed via IAP SSH "
        "tunneling (gcloud compute ssh --tunnel-through-iap). No VPN is "
        "required — IAP provides secure access."
    )
    add_diagram(doc,
        "Developer laptop --(IAP SSH tunnel)--> Dev VM (inside VPC) --> Cloud Run (internal)"
    )
    doc.add_paragraph(
        "When VPC-internal mode is selected, deploy.sh recommends enabling "
        "the dev VM since it is the primary access path. developer-setup.sh "
        "detects unreachable services and skips the smoke test with IAP-based "
        "remediation guidance."
    )

    doc.add_heading("Access Tiers (IAP-based, no VPN required)", level=2)
    doc.add_paragraph(
        "Both restricted-ingress modes use IAP as the access mechanism:"
    )
    add_table(doc,
        ["Tier", "Mode", "Access Path", "Best For"],
        [
            ["Tier 1", "GLB + IAP",
             "GLB fronts all services. IAP for browsers, token validation for APIs.",
             "Production deployments with custom domains"],
            ["Tier 2", "Dev VM + IAP SSH",
             "No GLB. Cloud Run uses --ingress internal. SSH to dev VM via IAP.",
             "Development/budget deployments"],
        ],
    )

    # =================================================================
    # PREREQUISITES AND DEPLOYMENT PATHS
    # =================================================================
    doc.add_page_break()
    doc.add_heading("Prerequisites and Deployment Paths", level=1)

    doc.add_heading("Prerequisites", level=2)
    doc.add_paragraph("Before you deploy, you need:")
    prereqs = [
        "A Google Cloud project where you have the Owner or Editor role.",
        "A billing account linked to that project. Running this costs roughly "
        "$0-5/month when idle.",
        "Access to the Anthropic Claude models on Vertex AI. Open the Vertex "
        "AI Model Garden, search for Claude, and enable the models you want "
        "(Opus 4.6, Sonnet 4.6, Haiku 4.5 are defaults).",
        "A local machine with gcloud CLI installed and logged in, and git installed. "
        "For the Terraform path: terraform >= 1.6. For the notebook path: just a browser.",
    ]
    for item in prereqs:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Four Ways to Deploy", level=2)
    doc.add_paragraph(
        "Pick whichever matches your comfort level. All four end up with the "
        "same resources."
    )

    add_table(doc,
        ["Path", "Method", "Best For"],
        [
            ["1. curl-to-bash",
             "curl -fsSL <deploy-url> | bash",
             "Fastest way to kick the tires; script clones repo into a temp dir"],
            ["2. Git clone + script (recommended)",
             "git clone <repo> && cd scripts && ./deploy.sh",
             "Keep a local copy you can inspect, modify, and re-run"],
            ["3. Terraform",
             "terraform init && terraform apply (two-phase)",
             "Teams that already use IaC; TF modules mirror config.yaml toggles"],
            ["4. Notebook (Colab / Vertex Workbench)",
             "Open deploy.ipynb and step through cells",
             "See each step and its output without installing anything locally"],
        ],
    )

    doc.add_heading("Cost Summary", level=2)
    doc.add_paragraph(
        "Idle is the common case. Typical monthly costs:"
    )
    add_table(doc,
        ["Configuration", "Estimated Cost"],
        [
            ["Default, idle (LLM + MCP + portal, no dev VM)", "~$0-5/month"],
            ["Default, light use (a few developers)", "~$10-30/month (Vertex tokens dominate)"],
            ["Everything on (incl. dev VM e2-small, BigQuery sink)", "~$25-50/month"],
            ["+ GLB (add to any config)", "+~$18/month"],
            ["+ VPC internal (add to any non-GLB config)", "+~$0 (VPC Connector is forced on)"],
        ],
    )
    doc.add_paragraph(
        "Token costs for Claude on Vertex follow Google's published Vertex AI "
        "pricing — they are billed to your GCP project, not to Anthropic."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 1
    # =================================================================
    doc.add_heading("1. The Orchestrator: deploy.sh", level=1)

    doc.add_paragraph(
        "deploy.sh is the single entry point for deploying the entire stack. "
        "It collects configuration interactively, writes a config.yaml file, "
        "and calls each component's deploy script in the correct order."
    )

    doc.add_heading("1.1 Bootstrap and Preflight", level=2)
    doc.add_paragraph(
        "The script first detects whether it is running from a git clone or "
        "via curl-to-bash. If piped from curl, it clones the repository to a "
        "temporary directory and re-executes itself from there. Then it:"
    )
    items = [
        "Sources lib/common.sh (strict mode, logging helpers, run_cmd for dry-run support) and lib/regions.sh (region picker).",
        "Runs require_cmd gcloud and require_cmd python3.",
        "Verifies an active gcloud account exists via gcloud auth list.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.2 Interactive Configuration Prompts", level=2)
    doc.add_paragraph(
        "The script prompts for each setting with sensible defaults:"
    )
    add_table(doc,
        ["Prompt", "Default", "Purpose"],
        [
            ["GCP Project ID", "gcloud config get-value project", "Every resource is created in this project"],
            ["Vertex region", "global", "Where Claude inference runs; global auto-routes to nearest backend"],
            ["Deploy LLM gateway?", "yes", "The core reverse proxy for Claude Code traffic"],
            ["Deploy MCP gateway?", "yes", "Shared MCP tool server"],
            ["Deploy dev portal?", "yes", "Self-service developer setup page"],
            ["Deploy dev VM?", "no", "Optional cloud development environment"],
            ["Install observability?", "yes", "BigQuery log sink and admin dashboard"],
            ["Deploy GLB?", "no", "Global Load Balancer for custom domains, IAP, Cloud Armor"],
            ["Restrict to VPC-internal?", "no (only if GLB=no)", "Cloud Run uses --ingress internal; developers access via IAP"],
            ["Allowed principals", "Current gcloud account", "Comma-separated user: and group: entries"],
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph(
        "If GLB is enabled, two follow-up prompts appear: (1) certificate mode — "
        "Google-managed cert with a custom domain, or self-signed for IP-only access; "
        "and (2) IAP support email (required for the OAuth consent screen when portal "
        "and GLB are both enabled)."
    )

    doc.add_paragraph(
        "If VPC-internal is enabled (mutually exclusive with GLB), the script "
        "displays guidance explaining that developers access services through the "
        "dev VM via IAP SSH tunneling — no VPN is required. If the dev VM was not "
        "already selected, the script recommends enabling it, since the dev VM is "
        "the primary access path in VPC-internal mode."
    )

    doc.add_heading("1.3 Config File and Confirmation Gate", level=2)
    doc.add_paragraph(
        "All answers are written to config.yaml at the repository root. The file "
        "includes project ID, region, component toggles, GLB settings, a vpc section "
        "(with internal_ingress flag), the allowed principals list, and pinned model "
        "versions. The file is displayed to the user, and the script asks: "
        "\"Proceed with this configuration?\" If declined, the config file is "
        "preserved for manual editing and the script exits."
    )

    doc.add_heading("1.4 GCP Setup", level=2)
    doc.add_paragraph("On confirmation, the script:")
    items = [
        "Sets the active gcloud project.",
        "Enables 11 required APIs in parallel: aiplatform, run, compute, iap, artifactregistry, cloudbuild, logging, monitoring, iamcredentials, secretmanager, serviceusage, bigquery.",
        "If Google-managed GLB cert was chosen, also enables dns.googleapis.com.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.5 Component Deployment Order", level=2)
    doc.add_paragraph(
        "Before deploying, the script validates mutual exclusion: GLB and "
        "VPC-internal cannot both be enabled. Components are then deployed "
        "conditionally, in a specific order:"
    )

    add_diagram(doc, """\
deploy.sh
  |-- Validate: GLB and VPC-internal are mutually exclusive
  |-- deploy-llm-gateway.sh      (always first; other components reference its URL)
  |-- deploy-mcp-gateway.sh
  |-- deploy-dev-portal.sh        (substitutes gateway URLs into HTML)
  |-- deploy-observability.sh     (BigQuery dataset + log sink + dashboard)
  |-- deploy-glb.sh               (optional, mutually exclusive with VPC-internal)
  |     +-- deploy-dev-portal.sh  (re-deployed to inject GLB URLs)
  +-- deploy-dev-vm.sh            (optional; always last; recommended for VPC-internal)""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The portal is intentionally deployed twice when GLB is enabled: first with "
        "Cloud Run URLs (so it exists for the GLB to reference), then again with GLB "
        "URLs (so developers see the correct endpoint). All deploy scripts use "
        "a three-way ingress conditional (standard / GLB / VPC-internal) to set "
        "the correct --ingress flag on each Cloud Run service."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 2
    # =================================================================
    doc.add_heading("2. The Core Gateway: deploy-llm-gateway.sh", level=1)
    doc.add_paragraph(
        "This script creates the LLM gateway — the most critical component in the stack. "
        "It is a thin FastAPI reverse proxy that sits between Claude Code and Vertex AI."
    )

    doc.add_heading("2.1 Service Account Creation", level=2)
    add_code_block(doc, """\
gcloud iam service-accounts create llm-gateway \\
  --display-name="LLM Gateway (Claude Code -> Vertex AI)" """)
    doc.add_paragraph("")
    doc.add_paragraph(
        "Idempotent: checks if the SA already exists before creating. After creation, "
        "calls wait_for_sa which polls for up to 15 seconds. GCP has eventual consistency "
        "between SA creation and IAM binding acceptance; without this pause, subsequent "
        "add-iam-policy-binding commands fail with INVALID_ARGUMENT."
    )

    doc.add_heading("2.2 IAM Role Grants", level=2)
    doc.add_paragraph("Three project-level roles are granted to the service account:")
    add_table(doc,
        ["Role", "Purpose"],
        [
            ["roles/aiplatform.user", "Call Vertex AI inference endpoints"],
            ["roles/logging.logWriter", "Write structured log entries to Cloud Logging"],
            ["roles/iam.serviceAccountViewer", "Resolve GCE service account emails from numeric IDs (needed for dev VM token validation)"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "This service account — not the developer — authenticates to Vertex AI. "
        "Developers never receive roles/aiplatform.user directly."
    )
    doc.add_paragraph(
        "The iam.serviceAccountViewer role is required because GCE metadata tokens "
        "(used by the dev VM) do not include an email in the tokeninfo response — "
        "only a numeric unique ID. The gateway resolves the email via the IAM API "
        "so it can match against ALLOWED_PRINCIPALS."
    )

    doc.add_heading("2.3 Container Image Build", level=2)
    doc.add_paragraph("The image is built via Cloud Build and pushed to Artifact Registry:")
    add_code_block(doc, """\
gcloud builds submit gateway/ --tag ${IMAGE} --quiet""")
    doc.add_paragraph("")
    doc.add_paragraph(
        "The Dockerfile is a two-stage build. The builder stage (python:3.12-slim) "
        "creates a virtual environment at /opt/venv and installs requirements.txt "
        "(fastapi, uvicorn, httpx, google-auth, google-cloud-logging). The runtime "
        "stage copies only the venv and application source, creates a non-root app "
        "user, and exposes port 8080. The CMD runs uvicorn with a single worker "
        "(Cloud Run scales horizontally by spawning container instances). "
        "Final image size is approximately 120 MB."
    )

    doc.add_heading("2.4 Cloud Run Deployment", level=2)
    doc.add_paragraph("The deploy command sets these flags:")
    add_code_block(doc, """\
gcloud run deploy llm-gateway \\
  --image ${IMAGE} \\
  --region ${FALLBACK_REGION} \\
  --service-account llm-gateway@${PROJECT_ID}.iam.gserviceaccount.com \\
  ${INGRESS_FLAG} \\
  --no-allow-unauthenticated \\
  --no-invoker-iam-check \\
  --min-instances 0 --max-instances 10 \\
  --cpu 1 --memory 512Mi \\
  --port 8080 \\
  --set-env-vars "GOOGLE_CLOUD_PROJECT=...;VERTEX_DEFAULT_REGION=...;\\
    ENABLE_TOKEN_VALIDATION=1;ALLOWED_PRINCIPALS=..."  """)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Why --no-invoker-iam-check: ")
    run.bold = True
    p.add_run(
        "Cloud Run's built-in IAM only accepts OIDC identity tokens. Claude Code "
        "sends OAuth2 access tokens from gcloud auth application-default login. "
        "These are fundamentally different token types — Cloud Run IAM rejects "
        "access tokens with 401. By disabling the platform IAM check, all requests "
        "reach the FastAPI container, where token_validation.py handles auth instead "
        "(it accepts both token types)."
    )

    p = doc.add_paragraph()
    run = p.add_run("Why --no-allow-unauthenticated: ")
    run.bold = True
    p.add_run(
        "Even though invoker IAM is disabled, this flag prevents the service from "
        "being marked as \"public\" in the GCP console, which some org policies audit."
    )

    doc.add_heading("2.5 Environment Variables", level=2)
    add_table(doc,
        ["Variable", "Value", "Purpose"],
        [
            ["GOOGLE_CLOUD_PROJECT", "Your project ID", "Logged in every request; used by google-auth for ADC"],
            ["VERTEX_DEFAULT_REGION", "Your region (e.g. global)", "Fallback when the URL path has no region"],
            ["ENABLE_TOKEN_VALIDATION", "1", "Registers the token validation middleware on startup"],
            ["ALLOWED_PRINCIPALS", "Comma-separated emails", "Only these identities can call the gateway"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "If ENABLE_VM=true, the dev VM's service account "
        "(claude-code-dev-vm@PROJECT.iam.gserviceaccount.com) is automatically "
        "appended to ALLOWED_PRINCIPALS."
    )

    doc.add_heading("2.6 Ingress Mode", level=2)
    doc.add_paragraph(
        "Each deploy script contains a three-way conditional that selects the "
        "ingress flag based on the deployment mode:"
    )
    add_table(doc,
        ["Mode", "Ingress Flag", "Behavior"],
        [
            ["Standard", "--ingress all", "Developer laptops reach the service directly; token validation is the security boundary"],
            ["GLB", "--ingress internal-and-cloud-load-balancing", "Only the GLB can reach the service; external direct access is blocked"],
            ["VPC-internal", "--ingress internal", "Only VPC traffic can reach the service; developers access via dev VM + IAP SSH"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "GLB and VPC-internal modes are mutually exclusive. In VPC-internal mode, "
        "the VPC Connector is forced on so Cloud Run egress routes through the VPC "
        "for Private Google Access. No VPN is required — IAP provides secure "
        "developer access."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 3
    # =================================================================
    doc.add_heading("3. Shared MCP Tools: deploy-mcp-gateway.sh", level=1)
    doc.add_paragraph(
        "Identical shape to the LLM gateway script, with these differences:"
    )
    add_table(doc,
        ["Aspect", "LLM Gateway", "MCP Gateway"],
        [
            ["Service account", "llm-gateway", "mcp-gateway"],
            ["IAM roles", "aiplatform.user + logging.logWriter + iam.serviceAccountViewer", "logging.logWriter + iam.serviceAccountViewer"],
            ["Source directory", "gateway/", "mcp-gateway/"],
            ["Framework", "FastAPI + httpx reverse proxy", "FastAPI + FastMCP (Streamable HTTP)"],
            ["Max instances", "10", "5"],
            ["MCP endpoint", "N/A", "/mcp"],
            ["Auth model", "--no-invoker-iam-check + token_validation.py", "Identical"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "The MCP gateway ships with one example tool (gcp_project_info) that returns "
        "project ID, number, region, and enabled API count. Tool-specific IAM roles "
        "are granted separately when a tool requires GCP API access. See "
        "ADD_YOUR_OWN_TOOL.md for the pattern."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 4
    # =================================================================
    doc.add_heading("4. Self-Service Portal: deploy-dev-portal.sh", level=1)
    doc.add_paragraph("This script builds and deploys a static HTML page for developer onboarding.")

    doc.add_heading("4.1 Placeholder Substitution", level=2)
    doc.add_paragraph(
        "Before building the container, the script substitutes four placeholders "
        "in dev-portal/public/index.html using sed:"
    )
    add_table(doc,
        ["Placeholder", "Replaced With"],
        [
            ["__LLM_GATEWAY_URL__", "Auto-discovered Cloud Run or GLB URL"],
            ["__MCP_GATEWAY_URL__", "Auto-discovered Cloud Run or GLB URL"],
            ["__PROJECT_ID__", "The GCP project ID"],
            ["__REGION__", "The Vertex AI region"],
        ],
    )

    doc.add_heading("4.2 Container and Deploy", level=2)
    doc.add_paragraph(
        "The container is nginx:1.27-alpine serving the substituted HTML. Port is "
        "set via Cloud Run's $PORT injection at container startup. Auth uses Cloud "
        "Run IAM (roles/run.invoker) per principal — browsers send OIDC tokens, "
        "which Cloud Run IAM accepts natively. The portal includes per-OS setup "
        "instructions (macOS, Linux, Windows via WSL2), a preview of settings.json, "
        "and a download link for developer-setup.sh."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 5
    # =================================================================
    doc.add_heading("5. Logging Pipeline: deploy-observability.sh", level=1)

    doc.add_heading("5.1 BigQuery Dataset", level=2)
    doc.add_paragraph(
        "The script creates a BigQuery dataset named claude_code_logs using the "
        "BigQuery REST API directly (not the bq CLI) for compatibility with "
        "enterprise HTTP proxies:"
    )
    add_code_block(doc, """\
POST https://bigquery.googleapis.com/bigquery/v2/projects/${PROJECT_ID}/datasets
Body: {
  "datasetReference": {"datasetId": "claude_code_logs"},
  "location": "${CR_REGION}",
  "description": "Claude Code gateway logs"
}""")

    doc.add_heading("5.2 Cloud Logging Sink", level=2)
    add_code_block(doc, """\
gcloud logging sinks create claude-code-gateway-logs \\
  "bigquery.googleapis.com/projects/${PROJECT_ID}/datasets/claude_code_logs" \\
  --log-filter='resource.type="cloud_run_revision" AND \\
    resource.labels.service_name=~"^(llm-gateway|mcp-gateway)$"' \\
  --use-partitioned-tables""")
    doc.add_paragraph("")
    doc.add_paragraph(
        "This captures only LLM and MCP gateway logs — not portal, dashboard, or dev VM logs. "
        "The sink auto-creates a partitioned BigQuery table named "
        "run_googleapis_com_stdout. The sink's writer identity (an auto-created "
        "service account) is granted roles/bigquery.dataEditor on the dataset."
    )

    doc.add_heading("5.3 Admin Dashboard", level=2)
    doc.add_paragraph(
        "A Python app deployed as a Cloud Run service. It dynamically discovers the "
        "BigQuery table by querying for tables matching the pattern run_googleapis_com_* "
        "in the claude_code_logs dataset (caches the result for 5 minutes). It exposes "
        "six API endpoints, each running a BigQuery SQL query:"
    )
    add_table(doc,
        ["Endpoint", "Chart Type", "Description"],
        [
            ["/api/requests-per-day", "Line chart", "Daily request volume, split by model (30-day window)"],
            ["/api/requests-by-model", "Bar chart", "Total requests per model"],
            ["/api/top-callers", "Bar chart", "Top 20 callers by request count"],
            ["/api/error-rate", "Time series", "Daily error percentage (status >= 400)"],
            ["/api/latency-percentiles", "Metric cards", "p50 / p95 / p99 latency to first header byte (7-day window)"],
            ["/api/recent-requests", "Data table", "Last 50 requests with caller, model, status, latency"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Query results are cached for 30 seconds. The frontend uses Chart.js and "
        "auto-refreshes every 60 seconds. Auth is via roles/run.invoker (same principal "
        "list as the gateways)."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 6
    # =================================================================
    doc.add_heading("6. Runtime Request Flow", level=1)
    doc.add_paragraph(
        "This section traces the path of a single Claude Code request through "
        "every component, from the developer's keystroke to the Vertex AI response."
    )

    doc.add_heading("6.1 FastAPI App Startup (main.py)", level=2)
    doc.add_paragraph("When the container starts, uvicorn loads app.main:app:")
    items = [
        "Lifespan hook creates a shared httpx.AsyncClient with timeouts: connect=10s, read=300s, write=60s, pool=10s. Stored in app.state.http_client. Connection pooling means all requests share TCP connections to Vertex.",
        "Token validation middleware is registered (because ENABLE_TOKEN_VALIDATION=1).",
        "Routes are registered in declaration order: GET /healthz (legacy), GET /health (liveness probe), then the catch-all /{full:path} route for proxying.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("6.2 Token Validation Middleware (token_validation.py)", level=2)
    doc.add_paragraph(
        "Every request passes through this middleware first. Health endpoints "
        "(/health, /healthz) are explicitly skipped."
    )
    doc.add_paragraph("The validation flow:")

    add_diagram(doc, """\
Request arrives
  |-- Path is /health or /healthz?
  |     +-- YES: Skip middleware, pass through to handler
  |
  |-- No Authorization header?
  |     +-- 401 {"error": "missing_token"}
  |
  |-- Not "Bearer <token>" format?
  |     +-- 401 {"error": "invalid_token"}
  |
  |-- Classify token type:
  |     |-- Starts with "ya29."? --> access token (skip JWT path)
  |     |-- 3 dot-separated parts, header >= 20 chars? --> JWT
  |     +-- Otherwise --> access token
  |
  |-- JWT path:
  |     +-- Verify via google.oauth2.id_token.verify_oauth2_token()
  |         (checks signature against Google's public keys)
  |         Extract email from claims
  |
  |-- Access token path:
  |     +-- Check TTL cache (1024 entries, 30-second TTL)
  |         |-- Cache hit: use cached email
  |         +-- Cache miss: call googleapis.com/tokeninfo
  |             |-- email in response? --> use it, cache result
  |             +-- No email (GCE metadata token)?
  |                   Resolve via IAM API: GET /v1/projects/-/serviceAccounts/{azp}
  |                   (azp = numeric SA unique ID from tokeninfo)
  |                   Cache resolved email (256 entries, 1-hour TTL)
  |
  |-- Verification failed (no email resolved)?
  |     +-- 401 {"error": "invalid_token", "detail": "..."}
  |
  |-- ALLOWED_PRINCIPALS set and email not in list?
  |     +-- 403 {"error": "forbidden", "detail": "X not in allowed list"}
  |
  +-- Store caller_email and caller_source in request.state
      Pass to next handler""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The ya29.* prefix check is critical: GCE metadata tokens (ya29.c.*) have "
        "three dot-separated parts, which would otherwise pass the JWT heuristic and "
        "be sent to OIDC verification, where they always fail. The prefix check "
        "short-circuits this misclassification."
    )
    doc.add_paragraph(
        "The IAM API fallback exists because GCE metadata tokens return only a numeric "
        "azp (authorized party) in the tokeninfo response, not an email. The gateway "
        "calls IAM to resolve the service account email, which is then matched against "
        "ALLOWED_PRINCIPALS. Results are cached for 1 hour to avoid per-request IAM "
        "API calls. This path requires roles/iam.serviceAccountViewer on the gateway SA."
    )

    doc.add_heading("6.3 Proxy Logic (proxy.py)", level=2)
    doc.add_paragraph(
        "Once past the middleware, the catch-all route calls proxy_request(). "
        "The function performs nine discrete steps:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 1 — Extract caller identity. ")
    run.bold = True
    p.add_run(
        "Checks request.state.caller_email (set by the token validation middleware) "
        "or falls back to the x-goog-authenticated-user-email header (set by Cloud "
        "Run or IAP). This email is logged but never forwarded to Vertex."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 2 — Extract model and region from URL path. ")
    run.bold = True
    p.add_run(
        "A regex extracts the model name (e.g. claude-sonnet-4-6) for logging. "
        "A second regex extracts the region (e.g. global or us-east5) for host routing."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 3 — Normalize path. ")
    run.bold = True
    p.add_run(
        "Claude Code omits the /v1/ prefix when ANTHROPIC_VERTEX_BASE_URL is set. "
        "Vertex AI requires it. The _normalize_path() function checks if the path "
        "starts with a version prefix — if not, it prepends /v1/."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 4 — Sanitize headers. ")
    run.bold = True
    p.add_run(
        "The sanitize_request_headers() function applies three tiers of header drops:"
    )

    add_table(doc,
        ["Tier", "Headers Dropped", "Reason"],
        [
            ["Exact match", "authorization, host, content-length, accept-encoding, x-cloud-trace-context, x-forwarded-*, forwarded", "Replaced with gateway values or not applicable upstream"],
            ["Hop-by-hop (RFC 7230)", "connection, keep-alive, proxy-authenticate, proxy-authorization, te, trailer, transfer-encoding, upgrade", "Must not be forwarded by proxies"],
            ["Prefix match", "anthropic-beta*, x-goog-*", "Vertex rejects unknown beta headers; x-goog-* are Cloud Run internals"],
        ],
    )
    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Step 5 — Acquire gateway SA token. ")
    run.bold = True
    p.add_run(
        "The get_vertex_access_token() function uses google.auth.default() with "
        "the cloud-platform scope to get the gateway SA's credentials, then refreshes "
        "the token if expired. Credentials are cached at module level with a threading "
        "lock. The call is wrapped in asyncio.to_thread() because google.auth is "
        "synchronous and would otherwise block the event loop."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 6 — Build upstream URL. ")
    run.bold = True
    p.add_run(
        "The region extracted in Step 2 maps to a hostname: \"global\" becomes "
        "aiplatform.googleapis.com; \"us-east5\" becomes "
        "us-east5-aiplatform.googleapis.com. The normalized path from Step 3 is "
        "appended to produce the full Vertex URL."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 7 — Forward request. ")
    run.bold = True
    p.add_run(
        "The full request body is read into memory (typically 4 KB to 800 KB of JSON). "
        "An httpx request is built with the cleaned headers and body, then sent with "
        "stream=True so the response body is consumed lazily."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 8 — Emit structured log. ")
    run.bold = True
    p.add_run(
        "One JSON line is emitted to stdout containing: caller email, caller source, "
        "HTTP method, request path, upstream host, Vertex region, model name, status "
        "code, latency to first header byte (ms), list of stripped beta headers, Cloud "
        "Run region, and project ID. Cloud Logging auto-parses this JSON and makes "
        "each field queryable as a jsonPayload.* field."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 9 — Stream response back. ")
    run.bold = True
    p.add_run(
        "A StreamingResponse wraps an async generator that yields raw chunks from "
        "Vertex via aiter_raw(). The response headers are filtered (transfer-encoding, "
        "connection, and content-length are removed since Starlette recomputes them). "
        "A BackgroundTask calls upstream.aclose() after the last byte is sent, "
        "releasing the connection back to the httpx pool."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 7
    # =================================================================
    doc.add_heading("7. Developer Setup: developer-setup.sh", level=1)
    doc.add_paragraph(
        "This script is run on each developer's machine after the infrastructure "
        "is deployed. It connects Claude Code to the gateway."
    )

    doc.add_heading("7.1 Prerequisites and Authentication", level=2)
    doc.add_paragraph(
        "The script requires gcloud and npm on PATH. It verifies an active gcloud "
        "account, then runs gcloud auth application-default login, which opens a "
        "browser and creates an OAuth2 refresh token at "
        "~/.config/gcloud/application_default_credentials.json. Claude Code reads "
        "this via ADC to get access tokens for each request."
    )

    doc.add_heading("7.2 Claude Code CLI Installation", level=2)
    add_code_block(doc, "npm install -g @anthropic-ai/claude-code")
    doc.add_paragraph("")
    doc.add_paragraph("Skipped if the claude command is already on PATH.")

    doc.add_heading("7.3 URL Discovery", level=2)
    doc.add_paragraph(
        "Gateway URLs are auto-discovered using a multi-source fallback chain:"
    )
    add_diagram(doc, """\
GLB domain (env var GLB_DOMAIN)
  +-- GLB static IP (gcloud compute addresses describe claude-code-glb-ip)
       +-- Cloud Run service URL (tries multiple regions):
            1. ${REGION} (from config)
            2. us-central1
            3. us-east5
            4. europe-west1""")
    doc.add_paragraph("")
    doc.add_paragraph(
        "For each region, the script runs gcloud run services describe llm-gateway "
        "and uses the first successful result. The same pattern is used for the MCP "
        "gateway. Discovered URLs are shown as defaults in the interactive prompt."
    )

    doc.add_heading("7.4 Settings File Generation", level=2)
    doc.add_paragraph(
        "The script uses Python to safely build JSON and writes to ~/.claude/settings.json:"
    )
    add_code_block(doc, """\
{
  "env": {
    "CLAUDE_CODE_USE_VERTEX": "1",
    "CLOUD_ML_REGION": "global",
    "ANTHROPIC_VERTEX_PROJECT_ID": "my-project-id",
    "ANTHROPIC_VERTEX_BASE_URL": "https://llm-gateway-xxx-uc.a.run.app",
    "CLAUDE_CODE_DISABLE_EXPERIMENTAL_BETAS": "1",
    "ANTHROPIC_DEFAULT_OPUS_MODEL": "claude-opus-4-6",
    "ANTHROPIC_DEFAULT_SONNET_MODEL": "claude-sonnet-4-6",
    "ANTHROPIC_DEFAULT_HAIKU_MODEL": "claude-haiku-4-5@20251001"
  },
  "mcpServers": {
    "gcp-tools": {
      "type": "http",
      "url": "https://mcp-gateway-xxx-uc.a.run.app/mcp"
    }
  }
}""")

    doc.add_heading("7.5 Environment Variable Reference", level=2)
    add_table(doc,
        ["Variable", "Effect on Claude Code"],
        [
            ["CLAUDE_CODE_USE_VERTEX=1", "Emit Vertex-format requests (not Anthropic API format)"],
            ["CLOUD_ML_REGION", "Embedded in URL path: /projects/P/locations/{REGION}/publishers/..."],
            ["ANTHROPIC_VERTEX_PROJECT_ID", "Embedded in URL path: /projects/{PROJECT_ID}/locations/..."],
            ["ANTHROPIC_VERTEX_BASE_URL", "Send requests here instead of aiplatform.googleapis.com (routes traffic through gateway)"],
            ["CLAUDE_CODE_DISABLE_EXPERIMENTAL_BETAS=1", "Suppress client-side anthropic-beta headers (server-side strip is belt-and-suspenders)"],
            ["ANTHROPIC_DEFAULT_*_MODEL", "Pin model versions to prevent unexpected behavior changes between releases"],
        ],
    )

    doc.add_heading("7.6 Reachability Pre-Check and Smoke Test", level=2)
    doc.add_paragraph(
        "Before the authenticated smoke test, the script probes the gateway's "
        "/healthz endpoint with a 5-second timeout. If the probe returns 000 "
        "(connection failure), the script recognizes the service likely uses "
        "internal-only ingress and skips the smoke test gracefully. It displays "
        "two IAP-based remediation options:"
    )
    items = [
        "If a GLB is deployed, re-run with the GLB URL (auto-discovered if available).",
        "SSH into the dev VM via IAP and run Claude Code from there.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")

    doc.add_paragraph("")
    doc.add_paragraph(
        "If the gateway is reachable, the script sends an ADC access token to "
        "/health — the same token type Claude Code sends. Results:"
    )
    add_table(doc,
        ["Status", "Meaning"],
        [
            ["200", "Smoke test passed — gateway is reachable and auth works"],
            ["401/403", "Identity may not be in the gateway's ALLOWED_PRINCIPALS list"],
            ["000", "Service uses internal-only ingress — access via GLB or dev VM"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "The script always exits 0 regardless of the smoke test result — the test "
        "is informational only. The final output prints \"setup complete\" with "
        "a clear summary: settings file path, smoke test status, and the command "
        "to start Claude Code."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 8
    # =================================================================
    doc.add_heading("8. MCP Gateway Runtime", level=1)

    doc.add_heading("8.1 Server Architecture (server.py)", level=2)
    doc.add_paragraph(
        "The MCP gateway is a FastAPI app with FastMCP mounted as a sub-application. "
        "Tools are registered via the @mcp.tool() decorator. The ASGI app is built "
        "in a version-adaptive way: it tries FastMCP 2.x's http_app() first, then "
        "falls back to the older streamable_http_app()."
    )
    doc.add_paragraph(
        "A critical design detail is the lifespan propagation. The parent FastAPI "
        "app's lifespan context manager manually forwards FastMCP's lifespan context. "
        "Without this, FastMCP's internal task group is never initialized and all tool "
        "calls fail with \"task group not initialized\" errors."
    )

    doc.add_heading("8.2 MCP Protocol Flow", level=2)
    doc.add_paragraph("When Claude Code invokes a tool, the protocol handshake is:")
    items = [
        "Client sends POST /mcp with JSON-RPC initialize request.",
        "Server responds with capabilities (tool list).",
        "Client sends notifications/initialized acknowledgment.",
        "Client sends tools/call with tool name and arguments.",
        "Server executes the tool function, returns result.",
        "Client sends DELETE /mcp to close the session.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Token validation is identical to the LLM gateway (the token_validation.py "
        "file is a copy; pre-deploy-check.sh verifies the two stay in sync)."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 9
    # =================================================================
    doc.add_heading("9. Populating the Dashboard: seed-demo-data.sh", level=1)
    doc.add_paragraph(
        "After a fresh deploy, the dashboard is empty. This script generates "
        "realistic traffic so the dashboard has data to display."
    )

    doc.add_heading("9.1 Configuration", level=2)
    add_table(doc,
        ["Parameter", "Default", "Flag"],
        [
            ["Users", "5", "--users N"],
            ["Requests per user", "10", "--requests-per-user N"],
            ["Duration", "30 minutes", "--duration-minutes N"],
            ["Total requests", "50 (computed)", "5 x 10"],
            ["Hard cap", "200", "--i-know-what-im-doing to exceed"],
        ],
    )

    doc.add_heading("9.2 Safety Controls", level=2)
    add_table(doc,
        ["Control", "Value"],
        [
            ["Model", "claude-haiku-4-5@20251001 (cheapest available)"],
            ["Max output tokens", "8 per request"],
            ["Estimated cost per request", "~$0.0001"],
            ["Cost for default run", "~$0.005 (50 requests)"],
            ["Confirmation prompt", "Shown when total exceeds 50 requests"],
            ["Hard cap", "200 requests (unless explicitly overridden)"],
        ],
    )

    doc.add_heading("9.3 Reachability Pre-Check", level=2)
    doc.add_paragraph(
        "Before sending any requests, the script probes the gateway's /healthz "
        "endpoint. If the probe returns 000 (connection failure), the script "
        "exits with an error and prints IAP SSH guidance: developers should SSH "
        "into the dev VM via IAP and run the script from there. This prevents "
        "wasting time on requests that will all fail in VPC-internal mode."
    )

    doc.add_heading("9.4 Request Execution", level=2)
    doc.add_paragraph(
        "The script cycles through a corpus of 20 developer-focused prompts "
        "(e.g. \"Refactor this function\", \"Explain this regex\", \"Write a "
        "unit test\"). Requests are spread evenly across the configured duration. "
        "Sleep between requests = duration_seconds / total_requests."
    )
    add_code_block(doc, """\
path="/v1/projects/${PROJECT_ID}/locations/${REGION}/publishers/anthropic/\\
  models/claude-haiku-4-5@20251001:rawPredict"

curl -X POST "${GATEWAY_URL}${path}" \\
  -H "Authorization: Bearer ${TOKEN}" \\
  -H "Content-Type: application/json" \\
  -d '{"anthropic_version":"vertex-2023-10-16",
       "messages":[{"role":"user","content":"..."}],
       "max_tokens":8}'""")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Attribution caveat: ")
    run.bold = True
    p.add_run(
        "All requests are attributed to the identity running the script. For "
        "multi-user demo data, run the script from different accounts."
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 10
    # =================================================================
    doc.add_heading("10. Viewing the Admin Dashboard", level=1)

    doc.add_heading("10.1 Accessing the Dashboard", level=2)
    doc.add_paragraph(
        "In standard mode, the dashboard is a Cloud Run service gated by "
        "roles/run.invoker. Access it by navigating to its URL with an identity "
        "token. In GLB mode, the dashboard is behind IAP — open the GLB URL in a "
        "browser and authenticate with your Google account."
    )
    add_code_block(doc, """\
DASHBOARD_URL=$(gcloud run services describe admin-dashboard \\
  --project $PROJECT_ID --region us-central1 \\
  --format="value(status.url)")

echo "${DASHBOARD_URL}"  # Open in browser""")

    doc.add_heading("10.2 Data Pipeline", level=2)
    doc.add_paragraph(
        "The end-to-end data flow from a developer request to a dashboard panel:"
    )
    add_diagram(doc, """\
1. Developer runs Claude Code and sends a prompt
2. Claude Code POSTs to the LLM gateway
3. Gateway proxies to Vertex AI, emits JSON log to stdout
4. Cloud Run captures stdout --> Cloud Logging
5. Log sink filter matches service_name --> BigQuery
6. BigQuery table receives a new row (partitioned by timestamp)
7. Dashboard queries BigQuery (30-second result cache)
8. Chart.js renders the data, auto-refreshes every 60 seconds""")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Typical latency from request to dashboard visibility is approximately "
        "60 seconds (Cloud Logging flush to BigQuery is the bottleneck)."
    )

    doc.add_heading("10.3 Dashboard Panels", level=2)
    add_table(doc,
        ["Panel", "Query Logic", "Time Window"],
        [
            ["Requests per Day", "COUNT(*) GROUP BY DATE(timestamp), model", "30 days"],
            ["Requests by Model", "COUNT(*) GROUP BY model ORDER BY count DESC", "30 days"],
            ["Top Callers", "COUNT(*) GROUP BY caller ORDER BY count DESC LIMIT 20", "30 days"],
            ["Error Rate", "COUNTIF(status_code >= 400) / COUNT(*) * 100 per day", "30 days"],
            ["Latency Percentiles", "APPROX_QUANTILES(latency_ms, 100) for p50, p95, p99", "7 days"],
            ["Recent Requests", "SELECT timestamp, caller, model, status, latency LIMIT 50", "All time"],
        ],
    )

    doc.add_heading("10.4 Expected Results After Seed Data", level=2)
    doc.add_paragraph(
        "With the default seed-demo-data.sh run (50 requests over 30 minutes), "
        "the dashboard will show:"
    )
    items = [
        "Requests per Day: one bar for today with approximately 50 requests, all Haiku.",
        "Requests by Model: single bar for claude-haiku-4-5@20251001.",
        "Top Callers: your email with 50 requests.",
        "Error Rate: 0% (assuming no errors).",
        "Latency Percentiles: p50 typically 200-500ms, p95 500-1500ms (Haiku is fast).",
        "Recent Requests: 50 entries, all 200 status, all Haiku.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # =================================================================
    # STAGE 11
    # =================================================================
    doc.add_heading("11. Access Tiers and Security Model", level=1)
    doc.add_paragraph(
        "The solution provides two IAP-based access tiers for restricted-ingress "
        "deployments. No VPN infrastructure is required — IAP handles secure "
        "developer access in both tiers."
    )

    doc.add_heading("11.1 Tier 1: GLB + IAP (Production)", level=2)
    doc.add_paragraph(
        "A Global HTTP(S) Load Balancer fronts all Cloud Run services. Cloud Run "
        "uses --ingress internal-and-cloud-load-balancing, meaning only the GLB "
        "can reach the services. Browser services (dev portal, admin dashboard) "
        "are protected by IAP, which handles Google SSO authentication. API services "
        "(llm-gateway, mcp-gateway) use app-level token validation — IAP is not "
        "applied to API backends because Claude Code sends Bearer tokens, not "
        "browser cookies."
    )
    add_diagram(doc, """\
Developer laptop
  |-- (HTTPS + access token) --> GLB --> llm-gateway (token_validation)
  |-- (HTTPS + access token) --> GLB --> mcp-gateway (token_validation)
  |-- (browser + IAP SSO)    --> GLB --> dev-portal  (IAP-protected)
  +-- (browser + IAP SSO)    --> GLB --> dashboard   (IAP-protected)""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "Requires a DNS domain for managed certificates. Self-signed certificates "
        "work with IP-only access but require NODE_TLS_REJECT_UNAUTHORIZED=0 in "
        "each developer's settings.json (set automatically by developer-setup.sh)."
    )

    doc.add_heading("11.2 Tier 2: Dev VM + IAP SSH (Development / Budget)", level=2)
    doc.add_paragraph(
        "No GLB needed. Cloud Run uses --ingress internal. Developers SSH into "
        "the dev VM via IAP TCP tunneling and run Claude Code directly from the "
        "VM. The dev VM is pre-configured with Claude Code, ADC credentials, and "
        "gateway URLs pointing at the internal Cloud Run services."
    )
    doc.add_paragraph(
        "deploy-dev-vm.sh handles two critical networking prerequisites "
        "automatically: (1) enabling Private Google Access on the default subnet "
        "(required for the no-public-IP VM to reach Cloud Run *.run.app endpoints), "
        "and (2) granting roles/run.invoker on the admin dashboard to the VM's "
        "service account (so developers can view the dashboard from the VM)."
    )
    add_diagram(doc, """\
Developer laptop
  |-- gcloud compute ssh --tunnel-through-iap claude-code-dev-shared
  |
  +-- Dev VM (inside VPC)
        |-- claude (Claude Code CLI, pre-installed)
        |-- settings.json points at internal Cloud Run URLs
        +-- Reaches llm-gateway, mcp-gateway directly (same VPC)""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "The dev VM has no public IP and is accessed exclusively via IAP. OS Login "
        "is enabled, so developers authenticate with their own Google identities. "
        "For developers who prefer their local editor, SSH port forwarding provides "
        "access to internal services:"
    )
    add_code_block(doc, """\
gcloud compute ssh claude-code-dev-shared --tunnel-through-iap \\
  --project=$PROJECT_ID --zone=$ZONE \\
  -- -L 8443:llm-gateway-abc123-uc.a.run.app:443""")

    doc.add_heading("11.3 Security Boundaries", level=2)
    add_table(doc,
        ["Layer", "Standard Mode", "GLB Mode (Tier 1)", "VPC-Internal (Tier 2)"],
        [
            ["Network ingress", "Public (--ingress all)", "GLB only", "VPC only"],
            ["API auth", "token_validation.py", "token_validation.py", "token_validation.py"],
            ["Browser auth", "Cloud Run IAM", "IAP (Google SSO)", "N/A (use dev VM)"],
            ["Developer access", "Direct from laptop", "Via GLB domain", "Via dev VM + IAP SSH"],
            ["VPN required?", "No", "No", "No"],
        ],
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 12 — Developer Onboarding
    # =================================================================
    doc.add_heading("12. Developer Onboarding", level=1)
    doc.add_paragraph(
        "This section covers the two workflows needed to get a new developer "
        "productive: the admin steps to grant access, and the developer steps "
        "to connect."
    )

    doc.add_heading("12.1 Admin: Onboarding a New Developer", level=2)
    doc.add_paragraph(
        "An admin (project owner/editor) performs these steps to grant a new "
        "developer access to the deployment:"
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 1 — Add the developer to ALLOWED_PRINCIPALS. ")
    run.bold = True
    p.add_run(
        "Update config.yaml (or the Terraform allowed_principals variable) to "
        "include the developer's identity, then re-deploy the gateways so the "
        "ALLOWED_PRINCIPALS environment variable is updated."
    )
    add_code_block(doc, """\
# Re-deploy gateways with updated principals
cd scripts && ./deploy.sh""")

    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("Step 2 — Grant IAP tunnel and OS Login roles (dev VM access). ")
    run.bold = True
    p.add_run(
        "If the developer will access the dev VM, grant two IAM roles:"
    )
    add_code_block(doc, """\
gcloud projects add-iam-policy-binding $PROJECT_ID \\
  --member="user:developer@example.com" \\
  --role="roles/iap.tunnelResourceAccessor" --condition=None --quiet

gcloud projects add-iam-policy-binding $PROJECT_ID \\
  --member="user:developer@example.com" \\
  --role="roles/compute.osLogin" --condition=None --quiet""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "These roles are granted automatically by deploy-dev-vm.sh for identities "
        "listed in PRINCIPALS at deploy time. For developers added later, the admin "
        "runs the commands above manually."
    )

    p = doc.add_paragraph()
    run = p.add_run("Step 3 (optional) — Grant dashboard access. ")
    run.bold = True
    p.add_run(
        "The admin dashboard uses Cloud Run IAM (roles/run.invoker). Grant it "
        "per-service:"
    )
    add_code_block(doc, """\
gcloud run services add-iam-policy-binding admin-dashboard \\
  --project=$PROJECT_ID --region=$REGION \\
  --member="user:developer@example.com" \\
  --role="roles/run.invoker" --quiet""")

    doc.add_heading("12.2 Developer: Connecting to the Dev VM", level=2)
    doc.add_paragraph(
        "Once the admin has completed onboarding, the developer connects with "
        "a single command:"
    )
    add_code_block(doc, """\
gcloud compute ssh claude-code-dev-shared \\
  --tunnel-through-iap \\
  --project=$PROJECT_ID \\
  --zone=$ZONE""")

    doc.add_paragraph("")
    doc.add_paragraph(
        "On first connection, gcloud generates an SSH key pair and pushes it "
        "via OS Login. This takes 10-20 seconds. Subsequent connections are "
        "instant."
    )
    doc.add_paragraph(
        "Once on the VM, the developer simply runs claude — the CLI is pre-installed "
        "and settings.json is pre-configured with the gateway URLs and project settings. "
        "No additional setup is required."
    )

    doc.add_heading("12.3 What the Dev VM Provides", level=2)
    add_table(doc,
        ["Component", "Details"],
        [
            ["Claude Code CLI", "Pre-installed via npm at startup"],
            ["settings.json", "Pre-configured with gateway URLs, project ID, region, model pins"],
            ["ADC credentials", "VM service account, auto-rotated by GCE metadata server"],
            ["VS Code Server", "Browser-based IDE (optional, controlled by install_vscode_server)"],
            ["Auto-shutdown", "Shuts down after N hours idle (default: 2) to save costs"],
        ],
    )

    doc.add_page_break()

    # =================================================================
    # STAGE 13
    # =================================================================
    doc.add_heading("13. Validation and Regression Testing", level=1)

    doc.add_heading("13.1 Pre-Deploy Checks: pre-deploy-check.sh", level=2)
    doc.add_paragraph(
        "A local code consistency validator that runs without GCP access. It "
        "validates cross-file parity across deploy scripts, Terraform modules, "
        "and application code. Currently runs 27 checks across six categories:"
    )
    add_table(doc,
        ["Category", "Checks", "What It Validates"],
        [
            ["Unit Tests", "1", "Gateway pytest suite passes"],
            ["Token Validation", "4", "Middleware in sync, registered conditionally, caller fallback, Dockerfile"],
            ["Deploy Script GLB", "7", "GLB conditionals, dev VM SA, orchestration order, certs, DNS, IAP"],
            ["Terraform Consistency", "6", "enable_glb vars, gateway_allowed_principals, backends, IAP, TLS"],
            ["VPC Internal Ingress", "4", "Three-way conditionals, enable_vpc_internal vars, reachability checks, exports"],
            ["IAP Access Model", "3", "No VPN in prompts, IAP guidance in scripts, IAP firewall rules"],
            ["Teardown Coverage", "1", "All GLB resources handled"],
        ],
    )

    doc.add_heading("13.2 End-to-End Tests: e2e-test.sh", level=2)
    doc.add_paragraph(
        "An 8-layer test suite that validates the full deployment from "
        "infrastructure through IAP access. Tests return PASS (0), FAIL (1), "
        "or SKIP (2, when a component is not deployed)."
    )
    add_table(doc,
        ["Layer", "Focus", "Example Tests"],
        [
            ["1. Infrastructure", "Cloud Run services, IP addresses", "Services in READY state, no public IPs"],
            ["2. Network Path", "Connectivity, auth", "Gateway /health, TLS, access token auth"],
            ["3. Gateway Proxy", "Inference, streaming", "Haiku inference, response streaming"],
            ["4. Dev Portal", "Static site", "Portal serves HTML, contains gateway URLs"],
            ["5. MCP Tools", "MCP protocol", "Tool list, tool invocation (gcp_project_info)"],
            ["6. Negative + Obs", "Security, dashboard", "Unauth rejected, dashboard health"],
            ["7. GLB", "Load balancer", "GLB health, inference, direct-run-blocked, unauth"],
            ["8. IAP Access", "IAP mechanisms", "SSH firewall rule, IAM bindings, dev VM reachability"],
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(
        "Layer 8 (IAP Access) was added to validate the IAP-based access model. "
        "Test 8.3 (dev VM reaches gateway internally) SSHs into the dev VM via "
        "IAP and curls the gateway from inside the VPC — confirming the full "
        "Tier 2 access path works end-to-end."
    )
    doc.add_paragraph(
        "If the gateway is unreachable (VPC-internal mode from outside the VPC), "
        "the script exits with code 2 and prints IAP SSH guidance instead of "
        "failing silently."
    )

    doc.add_page_break()

    # =================================================================
    # APPENDIX
    # =================================================================
    doc.add_heading("Appendix A: Complete Lifecycle Diagram", level=1)
    doc.add_paragraph(
        "The following diagram summarizes the entire deployment and operations "
        "workflow from initial deploy to dashboard visualization."
    )

    add_diagram(doc, """\
deploy.sh
  |
  |-- Prompts --> config.yaml (incl. vpc: internal_ingress flag)
  |-- Validates: GLB and VPC-internal are mutually exclusive
  |-- Enables APIs
  |
  |-- deploy-llm-gateway.sh
  |     |-- Create SA (llm-gateway)
  |     |-- Grant roles/aiplatform.user + logging.logWriter + iam.serviceAccountViewer
  |     |-- Docker build (multi-stage, ~120MB)
  |     |-- Three-way ingress: standard / GLB / VPC-internal
  |     +-- Cloud Run deploy (--no-invoker-iam-check, ENABLE_TOKEN_VALIDATION=1)
  |
  |-- deploy-mcp-gateway.sh  (same shape, + iam.serviceAccountViewer)
  |-- deploy-dev-portal.sh   (nginx, placeholder substitution, three-way ingress)
  |
  |-- deploy-observability.sh
  |     |-- BigQuery dataset (claude_code_logs)
  |     |-- Log sink (service_name = llm-gateway | mcp-gateway)
  |     +-- Admin dashboard (Cloud Run, Chart.js, 6 panels, three-way ingress)
  |
  |-- (optional) deploy-glb.sh   (mutually exclusive with VPC-internal)
  |     +-- deploy-dev-portal.sh (re-deployed to inject GLB URLs)
  +-- (optional) deploy-dev-vm.sh (recommended for VPC-internal; always last)
        |-- Enable Private Google Access on subnet
        |-- Grant VM SA roles/run.invoker on dashboard
        +-- Create GCE VM (no public IP, IAP SSH, OS Login)""")

    doc.add_paragraph("")

    add_diagram(doc, """\
developer-setup.sh (per developer laptop)
  |-- gcloud auth application-default login
  |-- npm install -g @anthropic-ai/claude-code
  |-- Auto-discover gateway URLs (GLB -> Cloud Run fallback chain)
  |-- Write ~/.claude/settings.json
  |-- Reachability pre-check (/healthz, 5s timeout)
  |     |-- Reachable:  authenticated smoke test (/health with ADC token)
  |     +-- Unreachable: skip test, print IAP SSH guidance (no VPN needed)
  +-- "setup complete" — start Claude Code with: claude""")

    doc.add_paragraph("")

    add_diagram(doc, """\
Access Tiers (both IAP-based, no VPN required)
  |
  |-- Tier 1: GLB + IAP (production)
  |     |-- Laptop --> GLB --> Cloud Run (internal-and-cloud-load-balancing)
  |     |-- Browser services: IAP (Google SSO)
  |     +-- API services: token_validation middleware
  |
  +-- Tier 2: Dev VM + IAP SSH (development / budget)
        |-- Laptop --> IAP SSH --> Dev VM (inside VPC)
        +-- Dev VM --> Cloud Run (internal) directly""")

    doc.add_paragraph("")

    add_diagram(doc, """\
claude (developer runs Claude Code)
  |-- Reads settings.json --> ANTHROPIC_VERTEX_BASE_URL
  |-- Sends POST /projects/P/locations/R/.../models/M:rawPredict
  |     (OAuth2 access token, no /v1/ prefix)
  |-- --> token_validation.py classifies token (ya29.* = access, else JWT)
  |       validates token, resolves email (IAM API fallback for GCE tokens)
  |       checks ALLOWED_PRINCIPALS
  |-- --> proxy.py normalizes path (/v1/), strips headers, adds SA token
  |-- --> Vertex AI processes request
  |-- --> Response streamed back through gateway
  +-- --> Structured log --> Cloud Logging --> BigQuery --> Dashboard""")

    doc.add_paragraph("")

    add_diagram(doc, """\
seed-demo-data.sh
  |-- Reachability pre-check (exit with IAP SSH guidance if unreachable)
  |-- 50 tiny Haiku requests (~$0.005 total)
  |-- Spread over 30 minutes
  +-- Populates all 6 dashboard panels""")

    # ============ FOOTER ============
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("---")
    run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Claude Code on Vertex AI — Reference Architecture v1.0\n"
        "Apache 2.0 License"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x86, 0x8B)

    return doc


if __name__ == "__main__":
    doc = build_document()
    out = "/home/david/Anthropic-Google-Co-Innovation/Claude-VAI.docx"
    doc.save(out)
    print(f"Saved to {out}")
